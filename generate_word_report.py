#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script để tạo file Word tổng hợp kết quả LAB 4
"""

import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Tạo document
doc = Document()

# ============ TIÊU ĐỀ =============
title = doc.add_heading('TÓNG HỢP KẾT QUẢ CÁC BÀI TẬP LAB 4', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Sử dụng NumPy để phân tích dữ liệu')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True

# ============ BÀI 1 =============
doc.add_heading('BÀI 1: PHÂN TÍCH ĐIỂM HỌC PHẦN (MIỀN GIÁO DỤC)', 1)

doc.add_heading('Input', 2)
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('10 sinh viên, 4 thành phần điểm (chuyên cần, giữa kỳ, thực hành, cuối kỳ)')
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('Trọng số: 10%, 20%, 30%, 40%')

doc.add_heading('Output Chính', 2)

# Tính toán bài 1
scores = np.array([
    [8.0, 7.5, 8.5, 7.0],
    [6.5, 6.0, 7.0, 6.5],
    [9.0, 8.5, 9.0, 8.5],
    [5.0, 5.5, 6.0, 5.5],
    [7.5, 7.0, 8.0, 7.5],
    [4.5, 5.0, 5.5, 5.0],
    [8.5, 9.0, 8.0, 9.0],
    [6.0, 6.5, 6.0, 6.5],
    [7.0, 7.5, 7.0, 8.0],
    [9.5, 9.0, 9.5, 9.0]
])
weights = np.array([0.1, 0.2, 0.3, 0.4])
final_score = scores @ weights

# Xếp loại
def classify_grade(score):
    if score >= 8.5:
        return 'A'
    elif score >= 7.0:
        return 'B'
    elif score >= 5.5:
        return 'C'
    elif score >= 4.0:
        return 'D'
    else:
        return 'F'

grades = np.array([classify_grade(score) for score in final_score])

p = doc.add_paragraph()
p.add_run(f'Shape: {scores.shape}  |  Ndim: {scores.ndim}  |  Dtype: {scores.dtype}').font.bold = True

p = doc.add_paragraph('\nĐiểm tổng kết (top 3):')
rank_idx = np.argsort(final_score)[::-1]
for i, idx in enumerate(rank_idx[:3]):
    p = doc.add_paragraph(f'  {i+1}. Sinh viên {idx+1}: {final_score[idx]:.2f} (Loại {grades[idx]})', style='List Number')

p = doc.add_paragraph('\nXếp loại:')
grade_counts = {grade: np.sum(grades == grade) for grade in ['A', 'B', 'C', 'D', 'F']}
for grade, count in sorted(grade_counts.items()):
    if count > 0:
        if grade == 'A':
            text = f'  Loại {grade} (≥8.5): {count} sinh viên'
        elif grade == 'B':
            text = f'  Loại {grade} (7.0-8.4): {count} sinh viên'
        elif grade == 'C':
            text = f'  Loại {grade} (5.5-6.9): {count} sinh viên'
        else:
            text = f'  Loại {grade} (4.0-5.4): {count} sinh viên'
        doc.add_paragraph(text)

# Sinh viên cần hỗ trợ
low_component = np.any(scores < 5.0, axis=1)
low_students = np.where(low_component)[0]
if len(low_students) > 0:
    p = doc.add_paragraph('\nSinh viên cần hỗ trợ:')
    for idx in low_students:
        components = []
        for j, comp in enumerate(['chuyên cần', 'giữa kỳ', 'thực hành', 'cuối kỳ']):
            if scores[idx][j] < 5.0:
                components.append(f'{comp} {scores[idx][j]}')
        p = doc.add_paragraph(f'  Sinh viên {idx+1}: {final_score[idx]:.2f} điểm ({", ".join(components)} < 5.0)', style='List Bullet')

# Z-score
z_scores = (scores[:, 3] - scores[:, 3].mean()) / scores[:, 3].std()
p = doc.add_paragraph('\nZ-score analysis (điểm cuối kỳ):')
p = doc.add_paragraph(f'  - Trung bình: {scores[:, 3].mean():.2f}')
p = doc.add_paragraph(f'  - Độ lệch chuẩn: {scores[:, 3].std():.2f}')
p = doc.add_paragraph(f'  - Phạm vi: {scores[:, 3].min():.1f} - {scores[:, 3].max():.1f}')

p = doc.add_paragraph('\nMức độ phân hóa:')
if scores[:, 3].std() < 1.0:
    comment = 'Mức độ phân hóa THẤP - điểm cuối kỳ tương đối đồng đều'
elif scores[:, 3].std() < 2.0:
    comment = 'Mức độ phân hóa TRUNG BÌNH - có sự khác biệt vừa phải'
else:
    comment = 'Mức độ phân hóa CAO - sự khác biệt rõ rệt'
doc.add_paragraph(f'  {comment}')

p = doc.add_paragraph('\nKỹ thuật chính:')
doc.add_paragraph('Matrix multiplication: scores @ weights', style='List Bullet')
doc.add_paragraph('Boolean indexing: scores < 5.0', style='List Bullet')
doc.add_paragraph('Z-score normalization: (x - mean) / std', style='List Bullet')

# ============ BÀI 2 =============
doc.add_page_break()
doc.add_heading('BÀI 2: PHÂN TÍCH CHUYÊN CẦN (MIỀN GIÁO DỤC)', 1)

doc.add_heading('Input', 2)
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('12 sinh viên, 8 buổi học')
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('Quy ước: 1 = có mặt, 0 = vắng')

doc.add_heading('Output Chính', 2)

# Tính toán bài 2
attendance = np.array([
    [1,1,1,1,1,1,1,1],
    [1,1,0,1,1,0,1,1],
    [1,0,0,1,1,1,0,1],
    [1,1,1,1,0,1,1,1],
    [0,1,1,0,1,1,1,0],
    [1,1,1,1,1,1,0,1],
    [1,0,1,0,1,0,1,0],
    [1,1,1,1,1,1,1,0],
    [0,0,1,1,0,1,1,1],
    [1,1,1,0,1,1,1,1],
    [1,1,0,0,1,0,1,1],
    [1,1,1,1,1,0,1,1]
])

present_count = attendance.sum(axis=1)
rate = present_count / attendance.shape[1] * 100
warning_idx = np.where(rate < 75)[0]
absent_count_by_session = (attendance == 0).sum(axis=0)
worst_session = np.argmax(absent_count_by_session)
full_attendance = np.where(np.all(attendance == 1, axis=1))[0]
two_absent_in_row = np.where(np.any((attendance[:, :-1] == 0) & (attendance[:, 1:] == 0), axis=1))[0]

p = doc.add_paragraph('Chuyên cần tổng quan:')
p = doc.add_paragraph(f'  - Tỉ lệ trung bình: {rate.mean():.1f}%')
p = doc.add_paragraph(f'  - Sinh viên đầy đủ: {len(full_attendance)} ({len(full_attendance)/len(attendance)*100:.0f}%)')
p = doc.add_paragraph(f'  - Sinh viên cảnh báo (<75%): {len(warning_idx)} ({len(warning_idx)/len(attendance)*100:.0f}%)')

if len(warning_idx) > 0:
    p = doc.add_paragraph('\nSinh viên bị cảnh báo:')
    for idx in warning_idx:
        p = doc.add_paragraph(f'  SV{idx+1}: {rate[idx]:.1f}% ({present_count[idx]}/8 buổi)', style='List Bullet')

p = doc.add_paragraph('\nBuổi học có vắng nhiều nhất:')
p = doc.add_paragraph(f'  Buổi {worst_session+1}: {absent_count_by_session[worst_session]} sinh viên vắng')

if len(two_absent_in_row) > 0:
    p = doc.add_paragraph('\nSinh viên có 2 buổi vắng liên tiếp:')
    for idx in two_absent_in_row:
        p = doc.add_paragraph(f'  SV{idx+1}', style='List Bullet')

p = doc.add_paragraph('\nNhận xét:')
if rate.mean() >= 85:
    comment = '✓ Ý thức học tập TỐT'
elif rate.mean() >= 75:
    comment = '⚠ Ý thức học tập CẦN CẢI THIỆN'
else:
    comment = '❌ Ý thức học tập KHÔNG TỐT'
doc.add_paragraph(f'  {comment}')

p = doc.add_paragraph('\nKỹ thuật chính:')
doc.add_paragraph('np.where(): Tìm sinh viên thỏa điều kiện', style='List Bullet')
doc.add_paragraph('np.any(): Kiểm tra 2 buổi vắng liên tiếp', style='List Bullet')
doc.add_paragraph('np.all(): Kiểm tra tất cả điều kiện', style='List Bullet')

# ============ BÀI 3 =============
doc.add_page_break()
doc.add_heading('BÀI 3: PHÂN TÍCH DOANH THU BÁN HÀNG (MIỀN KINH DOANH)', 1)

doc.add_heading('Input', 2)
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('5 sản phẩm, 7 ngày')

doc.add_heading('Output Chính', 2)

# Tính toán bài 3
sales = np.array([
    [120, 150, 130, 140, 160],
    [125, 145, 128, 142, 158],
    [130, 155, 135, 150, 162],
    [135, 160, 140, 152, 168],
    [140, 165, 145, 155, 170],
    [138, 162, 142, 153, 169],
    [145, 170, 150, 160, 175]
])

daily_total = sales.sum(axis=1)
product_total = sales.sum(axis=0)
product_mean = sales.mean(axis=0)
product_std = sales.std(axis=0)

p = doc.add_paragraph()
p.add_run(f'Tổng doanh thu tuần: {int(sales.sum())} triệu đồng').font.bold = True

p = doc.add_paragraph('\nDoanh thu theo sản phẩm (sắp xếp từ cao đến thấp):')
sorted_products = np.argsort(product_total)[::-1]
for rank, prod_idx in enumerate(sorted_products[:3]):
    p = doc.add_paragraph(f'  {rank+1}. SP{prod_idx+1}: {int(product_total[prod_idx])} triệu đồng', style='List Number')

best_day = np.argmax(daily_total)
best_product = np.argmax(product_total)
p = doc.add_paragraph(f'\nNgày tốt nhất: Ngày {best_day+1} với {int(daily_total[best_day])} triệu đồng')
p = doc.add_paragraph(f'Sản phẩm bán tốt nhất: SP{best_product+1} với {int(product_total[best_product])} triệu đồng')

p = doc.add_paragraph('\nĐộ ổn định (độ lệch chuẩn - thấp hơn = ổn định hơn):')
sorted_stability = np.argsort(product_std)
for i, prod_idx in enumerate(sorted_stability[:3]):
    p = doc.add_paragraph(f'  {i+1}. SP{prod_idx+1}: {product_std[prod_idx]:.2f}', style='List Number')

# Kịch bản tăng 8%
new_sales = sales.astype(float).copy()
new_sales[:, [1, 4]] *= 1.08
before_total = sales.sum()
after_total = new_sales.sum()
increase = after_total - before_total

p = doc.add_paragraph(f'\nKịch bản tăng 8% cho SP2 & SP5:')
p = doc.add_paragraph(f'  - Tổng trước: {int(before_total)} triệu đồng')
p = doc.add_paragraph(f'  - Tổng sau: {int(after_total)} triệu đồng')
p = doc.add_paragraph(f'  - Mức tăng: +{int(increase)} triệu đồng ({increase/before_total*100:.2f}%)')

p = doc.add_paragraph('\nKhuyến nghị:')
doc.add_paragraph(f'✓ Ưu tiên bán SP{best_product+1} (cao nhất + ổn định)', style='List Bullet')
stable_product = np.argmin(product_std)
doc.add_paragraph(f'✓ SP{stable_product+1} có doanh thu ổn định nhất', style='List Bullet')

p = doc.add_paragraph('\nKỹ thuật chính:')
doc.add_paragraph('np.std(): Tìm sản phẩm ổn định', style='List Bullet')
doc.add_paragraph('Boolean indexing: Lọc ngày trên trung bình', style='List Bullet')

# ============ BÀI 4 =============
doc.add_page_break()
doc.add_heading('BÀI 4: QUẢN LÝ TỒN KHO (MIỀN KINH DOANH)', 1)

doc.add_heading('Input', 2)
p = doc.add_paragraph()
p.add_run('• ').font.bold = True
p.add_run('10 mặt hàng với tồn kho hiện tại, mức tối thiểu, giá nhập')

doc.add_heading('Output Chính', 2)

# Tính toán bài 4
stock = np.array([35, 8, 12, 5, 40, 18, 7, 22, 9, 15])
min_stock = np.array([20, 15, 15, 10, 25, 20, 12, 18, 12, 15])
price = np.array([30, 25, 28, 22, 35, 20, 18, 24, 19, 21])
need_import = np.maximum(min_stock - stock, 0)
cost = need_import * price
total_cost = cost.sum()
limited_need = np.clip(need_import, 0, 20)
limited_total_cost = (limited_need * price).sum()

p = doc.add_paragraph('\nTrạng thái kho:')
enough = np.sum(need_import == 0)
shortage = np.sum(need_import > 0)
p = doc.add_paragraph(f'  - Mặt hàng đủ: {enough}')
p = doc.add_paragraph(f'  - Mặt hàng thiếu: {shortage}')
p = doc.add_paragraph(f'  - Tổng thiếu hụt: {int(need_import.sum())} đơn vị')

shortage_indices = np.where(need_import > 0)[0]
if len(shortage_indices) > 0:
    p = doc.add_paragraph('\nTop 3 mặt hàng thiếu nhiều nhất:')
    top3_shortage = np.argsort(need_import)[::-1][:3]
    for i, idx in enumerate(top3_shortage):
        if need_import[idx] > 0:
            p = doc.add_paragraph(f'  {i+1}. HH{idx+1}: Thiếu {int(need_import[idx])} đơn vị → Chi phí {int(cost[idx])} triệu đ', style='List Number')

p = doc.add_paragraph(f'\nChi phí nhập hàng:')
p = doc.add_paragraph(f'  - Tổng chi phí (không giới hạn): {int(total_cost)} triệu đồng')
p = doc.add_paragraph(f'  - Tổng chi phí (giới hạn 20 đơn vị): {int(limited_total_cost)} triệu đồng')
if total_cost > limited_total_cost:
    p = doc.add_paragraph(f'  - Tiết kiệm: {int(total_cost - limited_total_cost)} triệu đồng')

p = doc.add_paragraph('\nMức độ thiếu hụt:')
if need_import.sum() < 10:
    doc.add_paragraph('  ✓ Mức độ thiếu hụt THẤP - không quá nghiêm trọng')
elif need_import.sum() < 30:
    doc.add_paragraph('  ⚠ Mức độ thiếu hụt TRUNG BÌNH - cần ưu tiên nhập')
else:
    doc.add_paragraph('  ❌ Mức độ thiếu hụt NGHIÊM TRỌNG - phải nhập ngay')

p = doc.add_paragraph('\nKỹ thuật chính:')
doc.add_paragraph('np.maximum(): Tính lượng cần nhập', style='List Bullet')
doc.add_paragraph('np.clip(): Giới hạn tối đa 20 đơn vị', style='List Bullet')

# ============ BÀI 5 =============
doc.add_page_break()
doc.add_heading('BÀI 5: TỔNG HỢP VÀ SO SÁNH HAI MIỀN', 1)

doc.add_heading('Câu 1: Dữ liệu ở hai bài giống nhau về cấu trúc?', 2)
doc.add_paragraph('✓ Cả hai đều là ma trận 2D', style='List Bullet')
doc.add_paragraph('✓ Theo dõi nhiều đối tượng qua nhiều tiêu chí', style='List Bullet')
doc.add_paragraph('✓ Dùng dtype float/numeric', style='List Bullet')
doc.add_paragraph('✓ Tính toán chỉ số tổng hợp (weighting/summing)', style='List Bullet')

doc.add_heading('Câu 2: Phép toán NumPy nào được dùng ở cả hai?', 2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Phép toán'
header_cells[1].text = 'Bài 1'
header_cells[2].text = 'Bài 3'

operations = [
    ('.sum()', '✓', '✓'),
    ('.mean()', '✓', '✓'),
    ('.std()', '✓', '✓'),
    ('.max(), .min()', '✓', '✓'),
    ('np.where()', '✓', '✓'),
    ('@ (matrix mult)', '✓', ''),
]

for i, (op, bai1, bai3) in enumerate(operations):
    row = table.rows[i+1]
    row.cells[0].text = op
    row.cells[1].text = bai1
    row.cells[2].text = bai3

doc.add_heading('Câu 3: Bài nào dùng boolean indexing nhiều hơn?', 2)
doc.add_paragraph('→ Bài 1 (Giáo dục) dùng nhiều hơn', style='List Bullet').bold = True
p = doc.add_paragraph('VÌ SAO?')
doc.add_paragraph('Bài 1: Phân loại, xếp hạng, cảnh báo → nhiều điều kiện', style='List Bullet')
doc.add_paragraph('Bài 3: Chủ yếu aggregation (tổng, trung bình) → ít điều kiện', style='List Bullet')

doc.add_heading('Câu 4: Lợi ích của vectorization so với vòng lặp?', 2)
doc.add_paragraph('5 lợi ích chính:', style='List Number')
doc.add_paragraph('Hiệu suất: 10-100x nhanh hơn', style='List Bullet 2')
doc.add_paragraph('Code ngắn gọn: 1 dòng vs 10 dòng', style='List Bullet 2')
doc.add_paragraph('Bộ nhớ: Arrays lưu trữ liên tiếp, cache-friendly', style='List Bullet 2')
doc.add_paragraph('Tránh lỗi: Không cần lo off-by-one errors', style='List Bullet 2')
doc.add_paragraph('Linh hoạt: Hoạt động với tensor bất kỳ', style='List Bullet 2')

doc.add_heading('Câu 5: Khi tăng 100 lần, NumPy hỗ trợ gì?', 2)
doc.add_paragraph('8 Cơ chế scalability:', style='List Number')
doc.add_paragraph('Advanced Indexing → Xử lý 1M phần tử nhanh', style='List Bullet 2')
doc.add_paragraph('Cache-friendly operations → Tối ưu CPU', style='List Bullet 2')
doc.add_paragraph('Memory Mapping → Xử lý file > RAM', style='List Bullet 2')
doc.add_paragraph('Broadcasting → Tiết kiệm bộ nhớ', style='List Bullet 2')
doc.add_paragraph('DTYPE optimization → float32 vs float64', style='List Bullet 2')
doc.add_paragraph('Parallel Processing → Numba, GPU (CuPy), Dask', style='List Bullet 2')
doc.add_paragraph('Chunking → Chia dữ liệu lớn thành chunks', style='List Bullet 2')
doc.add_paragraph('Efficient Sorting → O(n log n) complexity', style='List Bullet 2')

# ============ KẾT LUẬN =============
doc.add_page_break()
doc.add_heading('KẾT LUẬN', 1)
doc.add_paragraph('NumPy tối ưu cho phân tích dữ liệu 2D', style='List Bullet')
doc.add_paragraph('Vectorization giúp tăng tốc 26x-100x lần', style='List Bullet')
doc.add_paragraph('Boolean indexing là kỹ năng then chốt trong phân loại', style='List Bullet')
doc.add_paragraph('Các phép toán chung (sum, mean, std, where) được dùng ở cả 2 miền', style='List Bullet')
doc.add_paragraph('Khi dữ liệu tăng 100x, NumPy vẫn xử lý hiệu quả', style='List Bullet')

# ============ LƯU FILE =============
doc.save('H:\\HocTap\\Chuyende3\\lab4_doanthingocanh\\LAB4_TONG_HOP_KET_QUA.docx')
print("✓ Đã tạo file Word: LAB4_TONG_HOP_KET_QUA.docx")
print("✓ File được lưu trong thư mục: H:\\HocTap\\Chuyende3\\lab4_doanthingocanh\\")

